from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.pipeline import PipelineState
from models.reports import FinalReport, CompanyReportSection

OUTPUTS_DIR = Path(__file__).parent.parent.parent / "tests" / "outputs"

# Palette couleurs
COLOR_TITLE = RGBColor(0x1A, 0x1A, 0x2E)    # bleu nuit
COLOR_SECTION = RGBColor(0x16, 0x21, 0x3E)  # bleu foncé
COLOR_ACCENT = RGBColor(0x0F, 0x3F, 0x7E)   # bleu accent
COLOR_BODY = RGBColor(0x2C, 0x2C, 0x2C)     # gris foncé


def _set_font(run, size_pt: int, bold: bool = False, color: RGBColor = None):
  run.font.size = Pt(size_pt)
  run.font.bold = bold
  if color:
      run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int):
  p = doc.add_heading(text, level=level)
  p.runs[0].font.color.rgb = COLOR_SECTION if level == 1 else COLOR_ACCENT
  return p


def _bullet(doc: Document, text: str):
  p = doc.add_paragraph(style="List Bullet")
  run = p.add_run(text)
  run.font.size = Pt(11)
  run.font.color.rgb = COLOR_BODY
  return p


def _add_cover(doc: Document, report: FinalReport):
  doc.add_paragraph()
  doc.add_paragraph()
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = p.add_run("TPLF ANALYSIS REPORT")
  _set_font(run, 24, bold=True, color=COLOR_TITLE)

  p2 = doc.add_paragraph()
  p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run2 = p2.add_run(f"Decision: {report.case_reference}")
  _set_font(run2, 14, color=COLOR_SECTION)

  p3 = doc.add_paragraph()
  p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run3 = p3.add_run(f"Generated on {report.generation_date}")
  _set_font(run3, 11, color=COLOR_BODY)

  doc.add_page_break()


def _add_exec_summary(doc: Document, report: FinalReport):
  _heading(doc, "Executive Summary", level=1)
  p = doc.add_paragraph(report.exec_summary)
  p.runs[0].font.size = Pt(11)
  p.runs[0].font.color.rgb = COLOR_BODY
  doc.add_paragraph()


def _add_decision_kpis(doc: Document, report: FinalReport):
  _heading(doc, "Judicial decision analysis", level=1)

  kpis = report.decision_kpis
  table = doc.add_table(rows=1, cols=2)
  table.style = "Table Grid"
  hdr = table.rows[0].cells
  hdr[0].text = "Criterion"
  hdr[1].text = "Value"
  for cell in hdr:
      for run in cell.paragraphs[0].runs:
          run.font.bold = True
          run.font.color.rgb = COLOR_SECTION

  rows = [
      ("Infraction type", kpis.infraction_type),
      ("Period", kpis.period),
      ("Geographic scope", kpis.geographic_scope),
      ("Affected markets", ", ".join(kpis.affected_markets)),
      ("Estimated overcharge", f"{kpis.estimated_overcharge_pct}%" if kpis.estimated_overcharge_pct else "N/A"),
      ("Estimated total harm", kpis.total_damage_estimate or "N/A"),
      ("Claimant profile", kpis.claimant_profile),
  ]
  for label, value in rows:
      row = table.add_row().cells
      row[0].text = label
      row[1].text = value

  doc.add_paragraph()
  _heading(doc, "Key legal findings", level=2)
  for finding in kpis.key_legal_findings:
      _bullet(doc, finding)
  doc.add_paragraph()


def _add_company_section(doc: Document, company: CompanyReportSection):
  _heading(doc, f"#{company.rank} — {company.company_name}", level=1)

  p_meta = doc.add_paragraph()
  run = p_meta.add_run(f"{company.ticker}  |  ISIN: {company.isin}  |  TPLF score: {company.tplf_score}/100")
  _set_font(run, 10, color=COLOR_BODY)

  doc.add_paragraph()
  _heading(doc, "TPLF strategic angles and legal grounds", level=2)
  for arg in company.key_arguments:
      _bullet(doc, arg)

  if company.significant_weaknesses:
      doc.add_paragraph()
      _heading(doc, "Material risks for the funder", level=2)
      for w in company.significant_weaknesses:
          _bullet(doc, w)

  doc.add_paragraph()
  _heading(doc, "TPLF synthesis and recommendation", level=2)
  p = doc.add_paragraph(company.conclusion)
  p.runs[0].font.size = Pt(11)
  p.runs[0].font.italic = True
  p.runs[0].font.color.rgb = COLOR_BODY

  doc.add_page_break()


def generate_docx(report: FinalReport, output_path: Path) -> str:
  doc = Document()

  # Marges
  for section in doc.sections:
      section.top_margin = Cm(2.5)
      section.bottom_margin = Cm(2.5)
      section.left_margin = Cm(3)
      section.right_margin = Cm(3)

  _add_cover(doc, report)
  _add_exec_summary(doc, report)
  _add_decision_kpis(doc, report)
  _add_company_section(doc, report.company_1)
  _add_company_section(doc, report.company_2)
  _add_company_section(doc, report.company_3)

  doc.save(str(output_path))
  return str(output_path)


async def document_generator(state: PipelineState) -> dict:
  report: FinalReport = state["report_content"]

  OUTPUTS_DIR.mkdir(exist_ok=True)
  timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
  output_path = OUTPUTS_DIR / f"rapport_tplf_{timestamp}.docx"

  path = generate_docx(report, output_path)
  return {"output_docx_path": path}